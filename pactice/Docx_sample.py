import docx
import os
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

os.getcwd()

doc = docx.Document()


# this a sample how to edit the portion in the previously defined format
def read_docx():
    doc1 = docx.Document(r"D:\Project\RD template.docx")
    for i in range(len(doc1.paragraphs)):
        print(doc1.paragraphs[i].text)
    # doc1.save("Sample.docx")


# This is fo rHeading Style 1
def heading_styl(text):
    """
    This Function Takes input and return the heading Formatted Output for the text in Heading 1
    :param text:
    :return:
    """

    heading = doc.add_paragraph(style="Heading 1")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    heading = heading.add_run(f"{text}")
    heading.font.name = "Poppins"
    heading.font.size = Pt(18)
    heading.font.color.rgb = RGBColor(0, 0, 0)


# This is for heading style 2
def heading_styl2(text):
    """
    This Function Takes input and return the heading Formatted Output for the text in Heading 2
    :param text:
    :return:
    """
    heading = doc.add_paragraph(style="Heading 2")
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    heading = heading.add_run(f"{text}")
    heading.bold = False
    heading.font.name = "Poppins"
    heading.font.size = Pt(16)
    heading.font.color.rgb = RGBColor(0, 0, 0)


# This for normal paragraphs
def normal_paragraph(text):
    """
    Return the value of text in a specific desired format for the hole file
    :param text:
    :return:
    """
    para = doc.add_paragraph(style="Normal")
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para = para.add_run(f"{text}")
    para.font.name = "Poppins"


def sub_headings(text):
    """
    This function wil return the subheadings for the given input string.
    :param text:
    :return:
    """
    para = doc.add_paragraph(style="Normal")
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p = para.add_run(f"{text}")
    p.bold = True
    p.font.name = "Poppins"
    p.font.size = Pt(12)


# This is to redefine the format and write a code to develop the format of our own
def format_test():
    heading_styl("Report Name")
    # Normal Para in poppins
    normal_paragraph("Global Military Gas Mask Market")

    # Add headings in bold
    heading_styl("Upcoming")
    # Normal Para in poppins
    normal_paragraph("No")
    # *************************************
    heading_styl("Segments")
    heading_styl2("Segment")
    normal_paragraph("Type")
    heading_styl2("Sub-segments")
    normal_paragraph("Full face gas mask, Half face gas mask")
    # ***************************************
    # This portion is repetitive could be created as a function
    heading_styl2("Segment")
    normal_paragraph("Application")
    heading_styl2("Sub-segments")
    normal_paragraph("Chemical Defense, Biological Defense, Nuclear Defense, Radiological Defense")
    # *******************************
    heading_styl("Images")
    normal_paragraph("Null")
    heading_styl("Sector")
    normal_paragraph("Industrials")
    heading_styl("Industry Group")
    normal_paragraph("Capital Goods")
    heading_styl("Industry")
    normal_paragraph("Machinery")
    heading_styl("Sub-Industry")
    normal_paragraph("Industrial Machinery")
    heading_styl("Market")
    normal_paragraph("Global Market")
    heading_styl("Role")
    normal_paragraph("Global Market")
    heading_styl("Country")
    normal_paragraph("Global")
    heading_styl("Report Data")
    normal_paragraph("Null")
    heading_styl("Product ID")
    normal_paragraph("Null")
    heading_styl("Download")
    normal_paragraph("14")
    heading_styl("Image Alt")
    normal_paragraph("Null")
    heading_styl("Report Slug")
    normal_paragraph("Null")
    heading_styl("Meta Title")
    normal_paragraph("Null")
    heading_styl("Meta Description")
    normal_paragraph("Null")
    heading_styl("Pages")
    normal_paragraph("170")
    heading_styl("Report Type")
    normal_paragraph("""
    Name
    
    Lite
    
    Price
    
    2000
    
    Name
    
    Medium
    
    Price
    
    3000
    
    Name
    
    Full
    
    Price
    
    4000
    """)
    # Methodologies
    heading_styl("Methodologies")

    normal_paragraph("For the global Military Gas Mask market, our research methodology involved a mixture of primary "
                     "and secondary data sources. Key steps involved in the research process are listed below:")

    para1 = doc.add_paragraph("1.", style="Normal")
    p = para1.add_run(" Information Procurement: ")
    p.bold = True
    p.font.name = "Poppins"
    para1.add_run("This stage involved the procurement of market data or related information via primary and "
                  "secondary sources. The various secondary sources used included various company websites, "
                  "annual reports, trade databases, and paid databases such as Hoover’s, Bloomberg Business, Factiva, "
                  "and Avention. Our team did 45 primary interactions Globally which included several stakeholders "
                  "such as manufacturers, customers, key opinion leaders, etc. Overall, information procurement was "
                  "one of the most extensive stages in our research process.").font.name = "Poppins"
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    para1 = doc.add_paragraph("2.", style="Normal")
    p = para1.add_run(" Information Analysis: ")
    p.bold = True
    p.font.name = "Poppins"
    para1.add_run("This step involved triangulation of data through bottom up and top-down approach to estimate and "
                  "validate the total size and future estimate of the global Military Gas Mask market.").font.name = "Poppins "
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    para1 = doc.add_paragraph("3.", style="Normal")
    p = para1.add_run(" Report Formulation: ")
    p.bold = True
    p.font.name = "Poppins"
    para1.add_run("The final step entailed the placement of data points at appropriate market spaces in an attempt to "
                  "deduce viable conclusions.").font.name = "Poppins"
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    para1 = doc.add_paragraph("4.", style="Normal")
    p = para1.add_run(" Validation & Publishing: ")
    p.bold = True
    p.font.name = "Poppins"
    para1.add_run("Validation is the most important step in the process. Validation & re-validation via an "
                  "intricately designed process helped us finalize data-points to be used for final calculations. The "
                  "final market estimates and forecasts were then aligned and sent to our panel of industry experts "
                  "for validation of data. Once the validation was done the report was sent to our Quality Assurance "
                  "team to ensure adherence to style guides, consistency & design.").font.name = "Poppins"
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Analyst report
    heading_styl("Analyst Support")
    para = doc.add_paragraph(style="Normal")
    p = para.add_run("Customization Options")
    p.bold = True
    p.font.name = "Poppins"
    p.font.size = Pt(12)
    normal_paragraph("With the given market data, our dedicated team of analysts can offer you with the following "
                     "customization options are available for the global Military Gas Mask market:")

    para1 = doc.add_paragraph(style="Normal")
    p = para1.add_run("Product Analysis: ")
    p.bold = True
    p.font.name = "Poppins"
    para1.add_run(
        "Product matrix, which offers a detailed comparison of the product portfolio of companies.").font.name = "Poppins"
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    para1 = doc.add_paragraph(style="Normal")
    p = para1.add_run("Regional Analysis: ")
    p.bold = True
    p.font.name = "Poppins"
    para1.add_run(
        "Further analysis of the global Military Gas Mask market for additional countries.").font.name = "Poppins"
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    para1 = doc.add_paragraph(style="Normal")
    p = para1.add_run("Competitive Analysis: ")
    p.bold = True
    p.font.name = "Poppins"
    para1.add_run("Detailed analysis and profiling of additional market players & comparative analysis of competitive "
                  "products.").font.name = "Poppins "
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    para1 = doc.add_paragraph(style="Normal")
    p = para1.add_run("Go to Market Strategy: ")
    p.bold = True
    p.font.name = "Poppins"
    para1.add_run(
        "Find the high growth channels to invest your marketing efforts and increase your customer base.").font.name = "Poppins "
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    para1 = doc.add_paragraph(style="Normal")
    p = para1.add_run("Innovation Mapping: ")
    p.bold = True
    p.font.name = "Poppins"
    para1.add_run("Identify racial solutions and innovation, connected to deep ecosystems of innovators, start-ups, "
                  "academics, and strategic partners.").font.name = "Poppins "
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    para1 = doc.add_paragraph(style="Normal")
    p = para1.add_run("Category Intelligence: ")
    p.bold = True
    p.font.name = "Poppins"
    para1.add_run("Customized intelligence that is relevant to their supply markets which will enable them to make "
                  "smarter sourcing decisions and improve their category management.").font.name = "Poppins "
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    para1 = doc.add_paragraph(style="Normal")
    p = para1.add_run("Public Company Transcript Analysis: ")
    p.bold = True
    p.font.name = "Poppins"
    para1.add_run(
        "To improve the investment performance by generating new alpha and making better-informed decisions.").font.name = "Poppins "
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    para1 = doc.add_paragraph(style="Normal")
    p = para1.add_run("Social Media Listening: ")
    p.bold = True
    p.font.name = "Poppins"
    para1.add_run("To analyse the conversations and trends happening not just around your brand, but around your "
                  "industry as a whole, and using those insights to make better marketing decisions.").font.name = \
        "Poppins "
    para1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Dynamic Input Started From here onwards
    # Market Insight
    heading_styl("Market Insights")
    normal_paragraph("The global Military Gas Mask market was valued at US$ 5.51 Billion in 2021, and it is expected "
                     "to reach a value of US$ 14.93 Billion by 2028, at a CAGR of 11.7% over the forecast period ("
                     "2022–2028). Military Gas Mask automate the procedures of storing and transporting items along "
                     "the supply chain. These robots are frequently used in storage facilities and warehouses because "
                     "they provide higher levels of uptime than manual labor. It also leads to the biggest "
                     "productivity increases and profitability for end users. The utilization of warehouses outfitted "
                     "with self-driving technologies is becoming more popular. It emphasizes the significance of "
                     "on-time delivery across logistical infrastructures. Furthermore, the increased need for "
                     "sophisticated supply chain operations and E-commerce fulfillment services is a major driver "
                     "driving the creation of logistics robotics. At the moment, the use of AI, machine learning, "
                     "and warehouse management software solutions is creating chances for increased demand for "
                     "Military Gas Mask across the globe. Furthermore, labor-intensive warehouse operations are "
                     "forcing retailers and facility owners to embrace automated logistic systems. This is further "
                     "fueling the growth of the Military Gas Mask market.")

    heading_styl("Segmental Analysis")
    normal_paragraph("The Global Military Gas Mask Market is segmented on the basis of type, application, industry, "
                     "and region. Based on type, the market is segmented into Automated Guided Vehicles, Autonomous "
                     "Mobile Robots, Robot Arms, and Others. Based on application, the market is segmented into "
                     "Palletizing & De-palletizing, Pick & Place, Transportation, and Others. Based on industry, "
                     "the market is segmented into E-commerce, Healthcare, Retail, Food & Beverage, Automotive, "
                     "and Others. Based on region, the global Military Gas Mask market is segmented into North "
                     "America, Europe, Asia-Pacific, South America, and MEA.")

    # Get subheadings
    sub_headings("Analysis by Type:")
    normal_paragraph("The market is divided into autonomous mobile robots, automated guided vehicles, robot arms, "
                     "and others. The autonomous guided vehicles industry is predicted to increase at a substantial "
                     "rate during the forecast period owing to its ability to move freely around warehouses and "
                     "factories. Furthermore, using sensor and guided-vision technology, AGVs can execute an optimal "
                     "path in real-time or follow a pre-programmed path. Goods are labeled with barcodes and RFIDs, "
                     "allowing AGVs to recognize the product and deliver it to its intended location. The autonomous "
                     "mobile robots, robot arms, and other (UAVs) segments are expected to grow at a rapid pace "
                     "throughout the projection period due to their ability to do repetitive work at a low cost.")

    heading_styl("Regional Insights")
    normal_paragraph("In terms of revenue, Asia Pacific generated USD 1.78 Billion in 2021. The increased emphasis on "
                     "modernization efforts to improve existing technology, particularly in South Korea and China, "
                     "would stimulate growth. North America, on the other hand, is expected to maintain its lead due "
                     "to increased expenditures in adopting automated warehousing systems, smart factories, "
                     "and industry 4.0.")

    # Market Drivers
    heading_styl("Market Dynamics")
    sub_headings("Driver")
    normal_paragraph("•	Order fulfillment in the e-commerce sector is a significant contributor to the growing "
                     "deployment of Military Gas Mask. To stay up with delivery timetables, businesses are being "
                     "forced to employ the automation process to do monotonous operations using robotic solutions due "
                     "to an increase in the number of online customers. Furthermore, the successful integration of "
                     "digital automation networks enables a real-time perspective of activities. As a result, "
                     "manufacturers, shipping companies, and other businesses may simply assess the condition of "
                     "operations. This function is critical for the E-commerce sector because they are primarily "
                     "concerned with meeting the needs of their customers on time. This also contributes to the "
                     "market's expansion.")

    sub_headings("Restraint")
    normal_paragraph("•	The Military Gas Mask industry is expanding as a result of its capacity to execute a variety "
                     "of tasks. These robots are outfitted with a variety of programming software and sensors that "
                     "allow them to do multiple tasks in a row. This is due to the high expenses associated with "
                     "acquiring and programming such robots, which slows market growth. Furthermore, underdeveloped "
                     "economies, where regional firms find it difficult to spend heavily, are conducting research and "
                     "development operations to keep up with the top robotics advancements. This could lead to "
                     "increased market growth constraints. Furthermore, small and medium-sized merchants cannot "
                     "afford to invest in robotic solutions, so they hire laborers to perform repetitive jobs in "
                     "warehouses and distribution centers. This is further stifling market growth")
    # Competitive Markets
    heading_styl("Competitive Landscape")
    normal_paragraph("The Military Gas Mask market is relatively fragmented, with a high level of competition. Few "
                     "large players, like ABB Ltd, Honeywell, and Kiva Systems, now control the market in terms of "
                     "market share. These industry leaders are extending their customer base across several areas, "
                     "and many corporations are creating strategic and collaborative initiatives with other start-up "
                     "enterprises to enhance their market share and profitability.")

    sub_headings("Top Player’s Company Profiles")
    doc.add_paragraph(style="List Bullet").add_run("ABB (Switzerland)").font.name = "Poppins"
    doc.add_paragraph(style="List Bullet").add_run("Toyota Industries Corporation (Japan)").font.name = "Poppins"
    doc.add_paragraph(style="List Bullet").add_run("Kawasaki Heavy Industries, Ltd. (Japan)").font.name = "Poppins"
    doc.add_paragraph(style="List Bullet").add_run("Fanuc Corporation (Japan)").font.name = "Poppins"
    doc.add_paragraph(style="List Bullet").add_run("Vecna Robotics (US)").font.name = "Poppins"
    doc.add_paragraph(style="List Bullet").add_run("KUKA AG (Augsburg, Germany)").font.name = "Poppins"
    doc.add_paragraph(style="List Bullet").add_run("Yaskawa America, Inc. (US)").font.name = "Poppins"
    doc.add_paragraph(style="List Bullet").add_run("ADematic (US)").font.name = "Poppins"
    doc.add_paragraph(style="List Bullet").add_run("Krones AG (Germany)").font.name = "Poppins"
    doc.add_paragraph(style="List Bullet").add_run("I.M.A. Industria Macchine Automatiche S.P.A. (Italy)").font.name = "Poppins"
    doc.add_paragraph(style="List Bullet").add_run("Columbia/Okura LLC (US)").font.name = "Poppins"
    doc.add_paragraph(style="List Bullet").add_run("IAM Robotics (US)").font.name = "Poppins"

    sub_headings("Recent Developments")
    normal_paragraph("•	In October 2021, Geek announced the release of the Robo Shuttle RS8-DA, an 8-meter-high "
                     "flexible arm robot. The new robot, which has the largest capacity in the business, will allow "
                     "customers to make the best use of their warehouses.")
    # kEy Market Trends
    heading_styl("Key Market Trends")
    normal_paragraph("•	The concept of warehouses outfitted with self-driving systems is gaining traction. Robots "
                     "oversee the majority of the operational workflow, which includes scanning, pick and place, "
                     "transportation, and other tasks. With the increasing number of autonomous warehouses, "
                     "the industry is likely to rise at a rapid pace in the future years. Furthermore, "
                     "logistic service providers are rapidly investing, which is projected to boost the industry "
                     "further. For example, in 2019, XPO Logistics, a logistics service provider based in the United "
                     "States, committed USD 550 million each year to integrate industry-leading innovations across "
                     "its operations in North America and Europe. Furthermore, the growth of the E-commerce business "
                     "and the increasing necessity of real-time inventories are two major factors driving global "
                     "market growth.")

    # Sky quest analysis
    heading_styl("SkyQuest Analysis")
    normal_paragraph("SkyQuest’s ABIRAW (Advanced Business Intelligence, Research & Analysis Wing) is our Business "
                     "Information Services team that Collects, Collates, Co-relates and Analyses the Data collected "
                     "by means of Primary Exploratory Research backed by the robust Secondary Desk research.")

    normal_paragraph("According to our analysis, the expanding number of online consumers is propelling the "
                     "e-commerce industry all around the world. As a result, retailers are working hard to meet the "
                     "delivery deadlines. To accomplish this, they are utilizing logistical robots to assist in the "
                     "completion of complex jobs. Companies in underdeveloped nations, on the other hand, "
                     "have a variety of hurdles in deploying innovative robots in their businesses, "
                     "which is accompanied by a large investment demand. This aspect may limit the growth of the "
                     "Military Gas Mask market in the approaching years. Key industry players are implementing "
                     "cutting-edge technology to acquire a competitive advantage. To compete with their opponents, "
                     "the majority of them are engaging in collaborative activities such as agreements, contracts, "
                     "and joint ventures with local enterprises. Vecna Robotics, for instance, is regarded as a "
                     "market leader in the logistic robots business. It has a strong supply chain. Simultaneously, "
                     "in order to increase sales, it has implemented robotic warehouse technologies.")

    heading_styl("What's Included")
    normal_paragraph("NA")
    # Save to file
    doc.save("test.docx")


# doc = docx.Document(r"D:\Project\Udemy_practice\Sample.docx")
# print(len(doc.paragraphs))
# doc.paragraphs[].text = "Market Name"

# for i in range(len(doc.paragraphs)):
#    print(doc.paragraphs[i].text, i)

format_test()
